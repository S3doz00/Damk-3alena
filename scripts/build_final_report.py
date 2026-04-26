"""Build the final GP2 Report docx from the humanized markdown source.

- Parses damk-3alena/GP2_Report_Full.md
- Replaces every `*[Figure N - ...]*` placeholder with the matching image + caption
- Converts markdown tables to native docx tables (Table N captions)
- Renders code blocks in monospace
- Builds the prefatory matter required by the GP2 Template:
    cover page → certification page (blank template) → abstract →
    table of contents → list of figures → list of tables → list of symbols
- Applies ASU style: Times New Roman 12pt body, A4 2.5cm margins, line spacing 1.5

Run:
    /Users/.../damk-3alena/scripts/.venv/bin/python build_final_report.py
"""

from __future__ import annotations

import os
import re
from pathlib import Path
from typing import List

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
MD_SRC = ROOT / "GP2_Report_Full.md"
FIG_DIR = ROOT / "figures"
LOGO = ROOT / "submission" / "logo.png"
OUT = ROOT / "GP2_Report_Damk3alena.docx"

ASU_RED = RGBColor(0x8D, 0x00, 0x29)
DARK = RGBColor(0x10, 0x10, 0x10)
GREY = RGBColor(0x66, 0x66, 0x66)

# ---------------------------------------------------------------------------
# Figure registry — exhaustive list of 35 figures
# ---------------------------------------------------------------------------

FIGURES = {
    1:  ("fig01_gantt.png",                          "Figure 1 — Gantt Chart for the GP2 Project Plan"),
    2:  ("fig02_pert.png",                           "Figure 2 — PERT Chart with Critical Path"),
    3:  ("fig03_agile_scrum.png",                    "Figure 3 — Agile / Scrum Framework (adapted to Damk 3alena)"),
    4:  ("fig04_usecase.png",                        "Figure 4 — Damk 3alena Use Case Diagram"),
    5:  ("fig05_activity_register.png",              "Figure 5 — Activity Diagram: Donor Registration with OTP"),
    6:  ("fig06_activity_book.png",                  "Figure 6 — Activity Diagram: Book Appointment"),
    7:  ("fig07_activity_forgot.png",                "Figure 7 — Activity Diagram: Forgot Password"),
    8:  ("fig08_activity_create_request.png",        "Figure 8 — Activity Diagram: Create Blood Request"),
    9:  ("fig09_seq_register.png",                   "Figure 9 — Sequence Diagram: Donor Registration"),
    10: ("fig10_seq_book.png",                       "Figure 10 — Sequence Diagram: Book Appointment"),
    11: ("fig11_seq_forecast.png",                   "Figure 11 — Sequence Diagram: AI Demand Forecast"),
    12: ("fig12_seq_shortage.png",                   "Figure 12 — Sequence Diagram: Shortage Detection"),
    13: ("fig13_class.png",                          "Figure 13 — Detailed Class Diagram"),
    14: ("fig14_er.png",                             "Figure 14 — Database ER Diagram"),
    15: ("fig15_supabase_schema.png",                "Figure 15 — Supabase Schema Overview (14 Core Tables)"),
    16: ("fig16_ai_forecast.png",                    "Figure 16 — AI Service: Sample Forecast Output"),
    17: ("fig17_ai_shortage.png",                    "Figure 17 — AI Service: Sample Shortage Alerts"),
    18: ("fig18_ai_donor_recs.png",                  "Figure 18 — AI Service: Sample Donor Recommendations"),
    19: ("fig19_register.png",                       "Figure 19 — Mobile: Donor Registration Screen (Arabic UI shown)"),
    20: ("fig20_otp.png",                            "Figure 20 — Mobile: Email OTP Verification Screen (Arabic UI shown)"),
    21: ("fig21_home.png",                           "Figure 21 — Mobile: Home Screen"),
    22: ("fig22_map.png",                            "Figure 22 — Mobile: Map Screen with Facility Cards"),
    23: ("fig23_booking.png",                        "Figure 23 — Mobile: Appointment Booking Screen"),
    24: ("fig24_qr.png",                             "Figure 24 — Mobile: QR-Code Appointment Ticket"),
    25: ("fig25_profile.png",                        "Figure 25 — Mobile: Profile Screen"),
    26: ("fig26_dashboard_sidebar_expanded.jpg",     "Figure 26 — Dashboard: Sidebar Expanded"),
    27: ("fig27_dashboard_sidebar_collapsed.png",    "Figure 27 — Dashboard: Sidebar Collapsed"),
    28: ("fig28_dashboard_home.jpg",                 "Figure 28 — Dashboard: Home Page with Charts"),
    29: ("fig29_dashboard_create_request.jpg",       "Figure 29 — Dashboard: Create Request Form"),
    30: ("fig30_dashboard_requests.jpg",             "Figure 30 — Dashboard: Requests Page"),
    31: ("fig31_dashboard_ai_forecast.jpg",          "Figure 31 — Dashboard: AI Insights — Forecast Chart"),
    32: ("fig32_dashboard_ai_shortage.jpg",          "Figure 32 — Dashboard: AI Insights — Shortage Alerts"),
    33: ("fig33_dashboard_ai_donors.jpg",            "Figure 33 — Dashboard: AI Insights — Donor Recommendations"),
    34: ("fig34_mobile_arabic.png",                  "Figure 34 — Mobile: Arabic RTL Interface"),
    35: ("fig35_dashboard_arabic_rtl.png",           "Figure 35 — Dashboard: Arabic RTL Interface"),
}

ACRONYMS = [
    ("AI",      "Artificial Intelligence"),
    ("API",     "Application Programming Interface"),
    ("ASU",     "Applied Science Private University"),
    ("CORS",    "Cross-Origin Resource Sharing"),
    ("DOB",     "Date of Birth"),
    ("EHR",     "Electronic Health Records"),
    ("ER",      "Entity Relationship"),
    ("FR",      "Functional Requirement"),
    ("GPS",     "Global Positioning System"),
    ("IT",      "Information Technology"),
    ("JWT",     "JSON Web Token"),
    ("LIS",     "Laboratory Information System"),
    ("NFR",     "Non-Functional Requirement"),
    ("OTP",     "One-Time Password"),
    ("PERT",    "Program Evaluation and Review Technique"),
    ("PO",      "Product Owner"),
    ("REST",    "Representational State Transfer"),
    ("RLS",     "Row-Level Security"),
    ("RPC",     "Remote Procedure Call"),
    ("RTL",     "Right-to-Left"),
    ("SDK",     "Software Development Kit"),
    ("UAT",     "User Acceptance Testing"),
    ("UC",      "Use Case"),
    ("UI",      "User Interface"),
    ("UML",     "Unified Modeling Language"),
    ("URL",     "Uniform Resource Locator"),
    ("UUID",    "Universally Unique Identifier"),
    ("UX",      "User Experience"),
]

# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------

def new_doc() -> Document:
    doc = Document()

    # A4 portrait, 2.5cm margins
    s = doc.sections[0]
    s.page_width  = Cm(21.0)
    s.page_height = Cm(29.7)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(s, m, Cm(2.5))

    # Body style: Times New Roman 12pt, 1.5 line spacing
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), 'Times New Roman')
    rfonts.set(qn('w:hAnsi'),   'Times New Roman')
    rfonts.set(qn('w:cs'),      'Times New Roman')
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after  = Pt(6)

    # Heading hierarchy
    for level, size, before, after in [(1, 16, 18, 12), (2, 14, 14, 8), (3, 13, 10, 6), (4, 12, 8, 4)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Times New Roman"
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = DARK
        h.paragraph_format.space_before = Pt(before)
        h.paragraph_format.space_after  = Pt(after)
        if level == 1:
            h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_centered(doc: Document, text: str, *, bold=False, size=12, color=DARK,
                 italic=False, space_after=6) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)


# ---------------------------------------------------------------------------
# Cover, certification, abstract, ToC scaffolding
# ---------------------------------------------------------------------------

def add_cover(doc: Document) -> None:
    doc.add_paragraph()
    add_centered(doc, "Applied Science Private University", bold=True, size=14)
    add_centered(doc, "Faculty of Information Technology",  bold=True, size=14)
    doc.add_paragraph()
    doc.add_paragraph()
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Cm(6))
    doc.add_paragraph()
    add_centered(doc, "Graduation Project (2) Report", bold=True, size=18, color=ASU_RED)
    doc.add_paragraph()
    add_centered(doc, "Damk 3alena", bold=True, size=24, color=ASU_RED)
    add_centered(doc, "AI-Driven Blood Donation Prediction and Matching System for Jordan",
                 italic=True, size=14)
    doc.add_paragraph()
    doc.add_paragraph()

    add_centered(doc, "Prepared by:", bold=True, size=12)
    for name, sid in [
        ("Saad Abulihyeh",     "202210883"),
        ("Sondos Abulihyeh",   "202210881"),
        ("Raghad Al-Basha",    "202210479"),
        ("Rawan Daoud",        "202210309"),
    ]:
        add_centered(doc, f"{name}    {sid}", size=12)

    doc.add_paragraph()
    add_centered(doc, "Supervised by:", bold=True, size=12)
    add_centered(doc, "Dr. Jaber Al-Widian", size=12)
    doc.add_paragraph()
    doc.add_paragraph()
    add_centered(doc, "June 2026", size=14, bold=True)
    doc.add_paragraph()
    add_centered(doc, "Copyright © 2025–2026 Applied Science Private University. All rights reserved.",
                 size=10, color=GREY, italic=True)
    add_page_break(doc)


def add_certification(doc: Document) -> None:
    add_centered(doc, "Certification", bold=True, size=18, space_after=18)
    body = (
        "This is to certify that the graduation project entitled "
        "\"Damk 3alena — AI-Driven Blood Donation Prediction and Matching System for Jordan\" "
        "has been carried out by the team listed on the cover page under our supervision, "
        "and that the work submitted is to the best of our knowledge their original work, "
        "developed during the academic year 2025–2026 in partial fulfilment of the requirements "
        "for the degree of Bachelor of Science in Information Technology at the "
        "Applied Science Private University, Faculty of Information Technology."
    )
    p = doc.add_paragraph(body)
    p.paragraph_format.space_after = Pt(36)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Signature blocks
    for label in ["Project Supervisor", "Department Head", "Faculty Dean"]:
        add_centered(doc, "_____________________________", size=12, space_after=2)
        add_centered(doc, label, size=11, italic=True, space_after=20)

    add_page_break(doc)


def add_abstract(doc: Document, abstract_text: str) -> None:
    add_centered(doc, "Abstract", bold=True, size=18, space_after=18)
    p = doc.add_paragraph(abstract_text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_page_break(doc)


def add_field_paragraph(doc: Document, instr_text: str, fallback: str) -> None:
    """Insert a Word field that Saad can update with right-click → Update Field."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    instr = OxmlElement('w:instrText')
    instr.text = ' ' + instr_text + ' '
    run._r.append(instr)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)

    # placeholder text shown until user updates field
    placeholder = OxmlElement('w:t')
    placeholder.text = fallback
    run._r.append(placeholder)

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar3)


def add_toc(doc: Document) -> None:
    add_centered(doc, "Table of Contents", bold=True, size=18, space_after=12)
    add_field_paragraph(doc, 'TOC \\o "1-3" \\h \\z \\u',
                        "[Right-click here → Update Field to populate Table of Contents]")
    add_page_break(doc)


def add_list_of_figures(doc: Document) -> None:
    add_centered(doc, "List of Figures", bold=True, size=18, space_after=12)
    for n in sorted(FIGURES.keys()):
        caption = FIGURES[n][1]
        p = doc.add_paragraph(caption)
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.space_after = Pt(2)
    add_page_break(doc)


def add_list_of_tables(doc: Document, table_titles: List[str]) -> None:
    add_centered(doc, "List of Tables", bold=True, size=18, space_after=12)
    for i, title in enumerate(table_titles, start=1):
        p = doc.add_paragraph(f"Table {i} — {title}")
        p.paragraph_format.space_after = Pt(2)
    add_page_break(doc)


def add_list_of_symbols(doc: Document) -> None:
    add_centered(doc, "List of Symbols and Abbreviations", bold=True, size=18, space_after=12)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for sym, defn in ACRONYMS:
        row = table.add_row().cells
        row[0].text = sym
        row[1].text = defn
        for run in row[0].paragraphs[0].runs:
            run.bold = True
    # widen second column
    for row in table.rows:
        row.cells[0].width = Cm(3.0)
        row.cells[1].width = Cm(13.5)
    add_page_break(doc)


# ---------------------------------------------------------------------------
# Markdown → docx parser (focused on this report's actual patterns)
# ---------------------------------------------------------------------------

INLINE_BOLD = re.compile(r'\*\*(.+?)\*\*')
INLINE_ITALIC = re.compile(r'(?<!\*)\*(?!\*)(.+?)\*(?!\*)')
INLINE_CODE = re.compile(r'`([^`]+)`')
FIG_PLACEHOLDER = re.compile(r'^\*\[(?:Figure\s+(\d+)|Gantt Chart - Figure\s+(\d+)|PERT Chart - Figure\s+(\d+))\b.*?\]\*\s*$')


def render_inline(paragraph, text: str, *, italic=False) -> None:
    """Add `text` to `paragraph` honoring **bold**, *italic*, and `code` markers."""
    # Tokenize: split on bold/italic/code
    pieces = []
    i = 0
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|(?<![*])\*[^*]+\*(?![*]))')
    for m in pattern.finditer(text):
        if m.start() > i:
            pieces.append(("plain", text[i:m.start()]))
        token = m.group(0)
        if token.startswith("**"):
            pieces.append(("bold", token[2:-2]))
        elif token.startswith("`"):
            pieces.append(("code", token[1:-1]))
        else:
            pieces.append(("italic", token[1:-1]))
        i = m.end()
    if i < len(text):
        pieces.append(("plain", text[i:]))
    for kind, content in pieces:
        run = paragraph.add_run(content)
        if italic:
            run.italic = True
        if kind == "bold":
            run.bold = True
        elif kind == "italic":
            run.italic = True
        elif kind == "code":
            run.font.name = "Courier New"
            run.font.size = Pt(10)


def add_figure(doc: Document, n: int) -> None:
    if n not in FIGURES:
        return
    fname, caption = FIGURES[n]
    path = FIG_DIR / fname
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    if path.exists():
        p.add_run().add_picture(str(path), width=Cm(15))
    else:
        run = p.add_run(f"[MISSING IMAGE: {fname}]")
        run.italic = True
        run.font.color.rgb = ASU_RED
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(10)
    cr.font.color.rgb = GREY
    cap.paragraph_format.space_after = Pt(12)


def parse_markdown(doc: Document, md_text: str) -> List[str]:
    """Convert markdown body to docx blocks. Returns list of table titles encountered."""
    lines = md_text.splitlines()
    table_titles: List[str] = []
    i = 0
    in_code = False
    code_buf: List[str] = []
    code_lang = ""
    pending_table_title: str | None = None

    while i < len(lines):
        line = lines[i]

        # Code fences
        if line.startswith("```"):
            if in_code:
                # close
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_after = Pt(8)
                shading = OxmlElement('w:shd')
                shading.set(qn('w:val'), 'clear')
                shading.set(qn('w:color'), 'auto')
                shading.set(qn('w:fill'), 'F4F4F4')
                p.paragraph_format.element.get_or_add_pPr().append(shading)
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Courier New"
                run.font.size = Pt(9.5)
                code_buf = []
                in_code = False
            else:
                in_code = True
                code_lang = line.strip("` ").strip()
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Skip horizontal rules
        if stripped == "---":
            i += 1
            continue

        # Figure placeholder line
        m_fig = FIG_PLACEHOLDER.match(stripped)
        if m_fig:
            n = next((g for g in m_fig.groups() if g is not None), None)
            if n:
                add_figure(doc, int(n))
            i += 1
            continue

        # Headings
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("## "):
            heading = stripped[3:].strip()
            # Treat top-level chapter / section starts as new pages
            is_top_level = heading.startswith("Chapter") or heading.startswith("References")
            if is_top_level:
                add_page_break(doc)
            if heading == "Contents":
                # The markdown contents block is replaced by the auto ToC up front; skip it.
                i += 1
                while i < len(lines) and not lines[i].startswith("---") and not lines[i].startswith("## "):
                    i += 1
                continue
            if heading == "Abstract":
                # Abstract is rendered up front; skip its body in the markdown
                i += 1
                while i < len(lines) and not lines[i].startswith("---") and not lines[i].startswith("## "):
                    i += 1
                continue
            # Chapters and References are H1; everything else stays H2
            doc.add_heading(heading, level=1 if is_top_level else 2)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=4)
            i += 1
            continue

        # Markdown table:  starts with | and second line has --- separators
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r'^\|\s*[:-]+\s*(\|\s*[:-]+\s*)+\|\s*$', lines[i+1]):
            # collect rows
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            # rows[0] = header, rows[1] = separator, rows[2:] = data
            header = rows[0]
            data = rows[2:]
            tbl = doc.add_table(rows=1 + len(data), cols=len(header))
            tbl.style = "Table Grid"
            for j, h in enumerate(header):
                cell = tbl.rows[0].cells[j]
                cell.text = ""
                run = cell.paragraphs[0].add_run(h)
                run.bold = True
                # light header shading
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F0E6D8')
                tcPr.append(shd)
            for ri, row in enumerate(data, start=1):
                for ci, val in enumerate(row):
                    cell = tbl.rows[ri].cells[ci]
                    cell.text = ""
                    render_inline(cell.paragraphs[0], val)
            # spacer
            doc.add_paragraph()
            # If we had a pending "Table N: Title" line, record it
            if pending_table_title:
                table_titles.append(pending_table_title)
                pending_table_title = None
            else:
                table_titles.append(f"(Untitled table at line {i})")
            continue

        # Detect table title line:  "**Table NN: ...**"
        m_tt = re.match(r'^\*\*Table\s+\d+\s*[:—-]\s*(.+?)\*\*\s*$', stripped)
        if m_tt:
            # render the title as bold paragraph and remember for LoT
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(stripped.strip("*"))
            run.bold = True
            pending_table_title = m_tt.group(1).strip()
            i += 1
            continue

        # Bullet list (- ...) or numbered (1. ...)
        if re.match(r'^\s*[-*]\s+', line):
            indent = len(line) - len(line.lstrip())
            content = re.sub(r'^\s*[-*]\s+', '', line)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.6 + indent * 0.4)
            render_inline(p, content)
            i += 1
            continue
        if re.match(r'^\s*\d+\.\s+', line):
            content = re.sub(r'^\s*\d+\.\s+', '', line)
            p = doc.add_paragraph(style="List Number")
            render_inline(p, content)
            i += 1
            continue

        # Blank line → just a small spacer (skip — paragraph spacing handles it)
        if not stripped:
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        render_inline(p, stripped)
        i += 1

    return table_titles


# ---------------------------------------------------------------------------
# Markdown extraction helpers
# ---------------------------------------------------------------------------

def extract_abstract(md: str) -> str:
    m = re.search(r'^## Abstract\s*\n+(.+?)(?=\n---|\n## )', md, re.DOTALL | re.MULTILINE)
    return m.group(1).strip() if m else ""


def strip_cover_block(md: str) -> str:
    """Drop the title-page block at the very top of the markdown so the parser
    doesn't double-emit it."""
    # The first '## Abstract' marks where parseable content begins; we render
    # cover/abstract/lists separately. Cut everything before '## Abstract'
    # except keep '## Abstract' onward (parser will skip Abstract body).
    idx = md.find("## Abstract")
    if idx < 0:
        return md
    # Begin at the start of '## Abstract'. The parser will then skip to '## Chapter 1'.
    return md[idx:]


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    md = MD_SRC.read_text(encoding="utf-8")
    abstract = extract_abstract(md)
    body_md = strip_cover_block(md)

    doc = new_doc()

    # Prefatory matter
    add_cover(doc)
    add_certification(doc)
    add_abstract(doc, abstract)
    add_toc(doc)
    add_list_of_figures(doc)
    # List of tables needs to know title list — we'll render content first then prepend
    # Easier: render content in a temp doc to collect titles, then build prefatory + content.

    # Pass 1: parse content into a fresh second doc to collect table titles
    tmp = new_doc()
    table_titles = parse_markdown(tmp, body_md)

    # Build List of Tables from collected titles
    add_list_of_tables(doc, table_titles)
    add_list_of_symbols(doc)

    # Pass 2: actually render content into the final doc
    parse_markdown(doc, body_md)

    doc.save(OUT)
    print(f"Wrote {OUT}  ({OUT.stat().st_size // 1024} KB)")
    print(f"  paragraphs:  {len(doc.paragraphs)}")
    print(f"  tables:      {len(doc.tables)}")
    img_count = len(doc.inline_shapes)
    print(f"  inline images: {img_count}")
    print(f"  table titles collected: {len(table_titles)}")


if __name__ == "__main__":
    main()
